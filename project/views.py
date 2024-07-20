from django.shortcuts import redirect, render
from django.http import HttpResponse
from django.contrib.auth.models import User
from django.contrib import messages
from django.contrib.auth import authenticate,login,logout
import os
from docx2pdf import convert
from PyPDF2 import PdfReader
from django.conf import settings
from django.core.files.storage import FileSystemStorage
from docx import Document
import pandas as pd
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import WordNetLemmatizer
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
lemmatizer = WordNetLemmatizer()
import gensim.downloader as api
import string
import numpy as np
import nltk
from io import BytesIO

stop_words = set(stopwords.words('english'))
word2vec_model = api.load("word2vec-google-news-300")
import ollama
# Create your views here.
def home(request):
    return render(request,'index.html')
def student(request):
    return render(request,"student.html")
def teacher(request):
    return render(request,"teacher.html")
def signup(request):
    if request.method == "POST":
        username = request.POST.get('username')
        fname = request.POST.get('fname')
        lname = request.POST.get('lname')
        email = request.POST.get('email')
        pass1 = request.POST.get('pass1')
        pass2 = request.POST.get('pass2')

        myuser = User.objects.create_user(username, email, pass1)
        myuser.first_name = fname
        myuser.last_name = lname
        myuser.save()

        messages.success(request, "Your account has been successfully created")
        
        return redirect("signup")

    return render(request, "signup.html")

def teachersignup(request):
    if request.method == "POST":
        username1 = request.POST.get('username')
        fname1 = request.POST.get('fname')
        lname1 = request.POST.get('lname')
        email1 = request.POST.get('email')
        passt1 = request.POST.get('pass1')
        passt2 = request.POST.get('pass2')

        myuser1 = User.objects.create_user(username1, email1, passt1)
        myuser1.first_name = fname1
        myuser1.last_name = lname1
        myuser1.save()

        messages.success(request, "Your account has been successfully created")
        
        return redirect("teachersignup")

    return render(request, "loginteacher.html")

def signin(request):
    if request.method=="POST":
        username=request.POST.get('username')
        pass1=request.POST.get('pass1')

        user=authenticate(username=username,password=pass1)
        if user is not None:
            login(request,user)
            fname=user.first_name
            return render(request,"student.html",{'fname':fname})
        else:
            messages.error(request,"add credintials")
            return redirect('signup')
    return render(request,"signup.html")

def teachersignin(request):
    if request.method=="POST":
        username1=request.POST.get('username')
        passt1=request.POST.get('pass1')

        user1=authenticate(username=username1,password=passt1)
        if user1 is not None:
            login(request,user1)
            fname1=user1.first_name
            return render(request,"teacher.html",{'fname':fname1})
        else:
            messages.error(request,"add credintials")
            return redirect('teachersignup')
    return render(request,"loginteacher.html")

def signout(request):
    logout(request)
    messages.success(request,"logout successfully")
    return redirect('home')

def read_assignment_paper(input_file):
    doc = Document(input_file)
    questions = []
    answers = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.endswith('?'):
            questions.append(text.lower()) 
            if i + 1 < len(doc.paragraphs):
                answer = doc.paragraphs[i+1].text.strip().lower()
                answers.append(answer)
            else:
                answers.append("")
    final(questions,answers)

def read_question_paper(input_file):
    doc = Document(input_file)
    questions = []
    answers = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.endswith('?'):
            questions.append(text.lower()) 
            if i + 1 < len(doc.paragraphs):
                answer = doc.paragraphs[i+1].text.strip().lower()
                answers.append(answer)
            else:
                answers.append("")
    database(questions,answers)

def database(questions , answers):
    global df
    lemmatized_questions = []
    for question in questions:
        question_tokens = word_tokenize(question)
        question_filtered = [lemmatizer.lemmatize(word, pos='v')
	for word in question_tokens if word.lower() not in stop_words and word.lower() not in string.punctuation]
        lemmatized_questions.append(" ".join(question_filtered))
    
    lemmatized_answers = []
    for answer in answers:
        answer_tokens = word_tokenize(answer)
        answer_filtered = [lemmatizer.lemmatize(word, pos='v') for word in answer_tokens if word.lower() not in stop_words and word.lower() not in string.punctuation]
        lemmatized_answers.append(" ".join(answer_filtered))
    question_embeddings = [[word2vec_model[word] for word in question.split() if word in word2vec_model] for question in lemmatized_questions]
    answer_embeddings = [[word2vec_model[word] for word in answer.split() if word in word2vec_model] for answer in lemmatized_answers]

    embedding_dim = 300  
    avg_question_embeddings = [sum(embedding) / len(embedding) if embedding else [0]*embedding_dim for embedding in question_embeddings]
    avg_answer_embeddings = [sum(embedding) / len(embedding) if embedding else [0]*embedding_dim for embedding in answer_embeddings]

    df = pd.DataFrame({"Question": questions, "Answer": answers, "Question_Embedding": avg_question_embeddings, "Answer_Embedding": avg_answer_embeddings})
  
def compare_questions(user_question):
    user_embedding = np.array(user_question).reshape(1, -1)
    similarity_scores = [cosine_similarity(user_embedding, [answer_embedding]) for answer_embedding in df['Question_Embedding']]
    max_similarity_index = np.argmax(similarity_scores)
    max_similarity_score = similarity_scores[max_similarity_index][0]
    print(max_similarity_score)
    most_similar_question = df['Question'][max_similarity_index]
    key_answer = df.loc[max_similarity_index, 'Answer']
    return most_similar_question, key_answer

def generate_embeddings(user_text):
    lemmatized_user = []
    user_tokens = word_tokenize(user_text)
    user_filtered = [lemmatizer.lemmatize(word, pos='v') for word in user_tokens if word.lower() not in stop_words and word.lower() not in string.punctuation]
    lemmatized_user.append(" ".join(user_filtered))
    user_embeddings = [[word2vec_model[word] for word in answer.split() if word in word2vec_model] for answer in lemmatized_user]
    avg_user_embeddings = [sum(embedding) / len(embedding) if embedding else [0]*embedding_dim for embedding in user_embeddings]
    return avg_user_embeddings

def compare_answers(db_answer, user_answer):
    user_embedding = np.array(user_answer).reshape(1, -1)
    db_embeddings = np.array(db_answer).reshape(1, -1)
    cosine_distance = cosine_similarity(user_embedding, db_embeddings).reshape(1,-1)
    return cosine_distance
    # similarity_scores = [cosine_similarity(user_embedding, [answer_embedding]) for answer_embedding in df['Answer_Embedding']]
    # max_similarity_index = np.argmax(similarity_scores)
    # max_similarity_score = similarity_scores[max_similarity_index][0]
    # print(max_similarity_score)
    # most_similar_answer = df['Answer'][max_similarity_index]
    # cosine_distance = cosine_similarity(user_embedding, df['Answer_Embedding'][max_similarity_index].reshape(1,-1))
    # return most_similar_answer, cosine_distance
marks = 0
feedback = ""
def calculate_marks_and_feedback(percentage_similarity_score, user_answer, key_answer):
    ollama.pull('llama2')
    if percentage_similarity_score >= 90:
        marks = 10
        feedback = 'Congrats on solving the question correctly!'
    elif percentage_similarity_score >= 80:
        marks = 9
        feedback = 'Congrats on solving the question correctly!'
    elif percentage_similarity_score >= 70:
        marks = 8
        response = ollama.chat(model='llama2', messages=[
        {
        'role': 'user',
        'content': f"""user answer = {user_answer}
                    The actual answer: {key_answer}

                    Give the user feedback in form of bullets based on his answer. Keep it brief.""",
        },
        ] )
        feedback = response['message']['content']
    elif percentage_similarity_score >= 60:
        marks = 7
        response = ollama.chat(model='llama2', messages=[
        {
        'role': 'user',
        'content': f"""user answer = {user_answer}
                    The actual answer: {key_answer}

                    Give the user feedback in form of bullets based on his answer. Keep it brief.""",
        },
        ] )
        feedback = response['message']['content']
    elif percentage_similarity_score >= 50:
        marks = 6
        response = ollama.chat(model='llama2', messages=[
        {
        'role': 'user',
        'content': f"""user answer = {user_answer}
                    The actual answer: {key_answer}

                    Give the user feedback in form of bullets based on his answer. Keep it brief.""",
        },
        ] )
        feedback = response['message']['content']
    elif percentage_similarity_score >= 40:
        marks = 5
        response = ollama.chat(model='llama2', messages=[
        {
        'role': 'user',
        'content': f"""user answer = {user_answer}
                    The actual answer: {key_answer}

                    Give the user feedback in form of bullets based on his answer. Tell them their performance reached not too bad, it was average and what they can do to improve it. Keep it brief.""",
        },
        ] )
        feedback = response['message']['content']
    elif percentage_similarity_score >= 30:
        marks = 4
        response = ollama.chat(model='llama2', messages=[
        {
        'role': 'user',
        'content': f"""user answer = {user_answer}
                    The actual answer: {key_answer}

                    Give the user feedback in form of bullets based on his answer. Why his answer was wrong and what he can do to improve it. Keep it brief.""",
        },
        ] )
        feedback = response['message']['content']
    elif percentage_similarity_score >= 20:
        marks = 3
        response = ollama.chat(model='llama2', messages=[
        {
        'role': 'user',
        'content': f"""user answer = {user_answer}
                    The actual answer: {key_answer}

                    Give the user feedback in form of bullets based on his answer. Tell them their effort was not up to the mark and why their answer was wrong and what he can do to improve it. Keep it brief.""",
        },
        ] )
        feedback = response['message']['content']
    elif percentage_similarity_score >= 10:
        marks = 2
        response = ollama.chat(model='llama2', messages=[
        {
        'role': 'user',
        'content': f"""user answer = {user_answer}
                    The actual answer: {key_answer}

                    Give the user feedback in form of bullets based on his answer. Tell them their effort was not up to the mark and why their answer was wrong and what he can do to improve it. Keep it brief.""",
        },
        ] )
        feedback = response['message']['content']
    else:
        marks = 1
        response = ollama.chat(model='llama2', messages=[
        {
        'role': 'user',
        'content': f"""user answer = {user_answer}
                    The actual answer: {key_answer}

                    Give the user feedback in form of bullets based on his answer. Tell them their effort was not up to the mark and why their answer was wrong and what he can do to improve it. Keep it brief.""",
        },
        ] )
        feedback = response['message']['content']
    return marks, feedback 

def final(questions,user_answers):
    result_data = {'Question': [], 'Marks': [], 'Feedback': []}
    total_marks = 0

    for question, user_answer in zip(questions, user_answers):
        avg_embedding_ques = generate_embeddings(question)
        similar_question, key_answer = compare_questions(avg_embedding_ques)
        avg_user_ans_embeddings = generate_embeddings(user_answer)
        key_answer_embeddings = generate_embeddings(key_answer)

        cosine_answer = compare_answers(key_answer_embeddings, avg_user_ans_embeddings)
        percentage_similarity_score = cosine_answer * 100

        marks, feedback = calculate_marks_and_feedback(percentage_similarity_score, user_answer, key_answer)
        total_marks += marks

        result_data['Question'].append(question)
        result_data['Marks'].append(marks)
        result_data['Feedback'].append(feedback)

    # Create a DataFrame from the result data
    df = pd.DataFrame(result_data)

    # Save DataFrame to BytesIO object
    df.to_excel('summary.xlsx', index=False,engine="openpyxl")
    

def studentdocument(request):
    if request.method == 'POST' and request.FILES['document']:
        uploaded_file = request.FILES['document']
        fs = FileSystemStorage()
        filename = fs.save(uploaded_file.name, uploaded_file)
        uploaded_file_path = os.path.join(settings.MEDIA_ROOT, filename)

        # Call your conversion function
        read_assignment_paper(uploaded_file_path)

        return render(request, "student.html")

    return render(request, "student.html")

def teacherdocument(request):
    if request.method == 'POST' and request.FILES['document']:
        uploaded_file = request.FILES['document']
        fs = FileSystemStorage()
        filename = fs.save(uploaded_file.name, uploaded_file)
        uploaded_file_path = os.path.join(settings.MEDIA_ROOT, filename)

        # Call your conversion function
        read_question_paper(uploaded_file_path)

        return render(request, "teacher.html")

    return render(request, "teacher.html")